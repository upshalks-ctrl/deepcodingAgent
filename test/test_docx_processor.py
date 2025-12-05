#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX处理器测试
测试DocxProcessor的功能
"""

import asyncio
import sys
from pathlib import Path

# 设置标准输出编码为UTF-8
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.document_processors.docx_processor import DocxProcessor


class TestDocxProcessor:
    """DOCX处理器测试类"""

    def __init__(self):
        self.test_dir = Path(__file__).parent / "test_data"
        self.test_dir.mkdir(exist_ok=True)
        self.test_docx = self.test_dir / "test_document.docx"

    async def create_test_docx(self):
        """创建测试DOCX文件"""
        if self.test_docx.exists():
            return

        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()

            # 添加标题
            title = doc.add_heading('DOCX Processor Test Document', 0)

            # 添加段落
            doc.add_paragraph('This is a test document for DOCX processing.')

            # 添加章节
            doc.add_heading('Chapter 1: Introduction', level=1)

            doc.add_paragraph(
                'This chapter introduces the DOCX processor and its capabilities. '
                'The processor can extract text, paragraphs, and tables from Word documents.'
            )

            doc.add_heading('Section 1.1: Features', level=2)

            # 添加项目符号列表
            features = doc.add_paragraph()
            features.add_run('• Automatic text extraction\n')
            features.add_run('• Paragraph-based segmentation\n')
            features.add_run('• Table processing\n')
            features.add_run('• Metadata extraction')

            doc.add_heading('Section 1.2: Supported Formats', level=2)

            doc.add_paragraph(
                'The processor supports the following formats:\n'
                '1. DOCX (Microsoft Word 2007+)\n'
                '2. DOC (Legacy Word format via conversion)\n'
                '3. RTF (Rich Text Format)'
            )

            # 添加表格
            doc.add_heading('Chapter 2: Technical Details', level=1)

            table = doc.add_table(rows=4, cols=3)
            table.style = 'Table Grid'

            # 表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Feature'
            hdr_cells[1].text = 'Status'
            hdr_cells[2].text = 'Notes'

            # 数据行
            row_cells = table.rows[1].cells
            row_cells[0].text = 'Text Extraction'
            row_cells[1].text = 'Supported'
            row_cells[2].text = 'Complete'

            row_cells = table.rows[2].cells
            row_cells[0].text = 'Table Processing'
            row_cells[1].text = 'Supported'
            row_cells[2].text = 'Partial'

            row_cells = table.rows[3].cells
            row_cells[0].text = 'Image Extraction'
            row_cells[1].text = 'Not Supported'
            row_cells[2].text = 'Future version'

            doc.add_heading('Chapter 3: Performance', level=1)

            doc.add_paragraph(
                'Performance metrics for the DOCX processor:\n\n'
                '• Small documents (< 1MB): < 1 second\n'
                '• Medium documents (1-10MB): 1-5 seconds\n'
                '• Large documents (> 10MB): 5-30 seconds'
            )

            # 添加编号列表
            doc.add_heading('Implementation Steps', level=2)

            step1 = doc.add_paragraph('Initialize the DocxProcessor', style='List Number')
            step2 = doc.add_paragraph('Open the DOCX file', style='List Number')
            step3 = doc.add_paragraph('Extract text and metadata', style='List Number')
            step4 = doc.add_paragraph('Return Document objects', style='List Number')

            doc.add_heading('Conclusion', level=1)

            doc.add_paragraph(
                'This document demonstrates the capabilities of the DOCX processor. '
                'It can handle various document elements including paragraphs, headings, '
                'tables, and lists.'
            )

            doc.save(str(self.test_docx))
            print(f"✓ Created test DOCX: {self.test_docx}")
        except Exception as e:
            print(f"✗ Failed to create test DOCX: {e}")
            raise

    async def test_python_docx_backend(self):
        """测试python-docx后端"""
        print("\n" + "="*70)
        print("TEST 1: python-docx Backend")
        print("="*70)

        processor = DocxProcessor()

        try:
            documents = await processor.process(str(self.test_docx), backend='python-docx')
            print(f"✓ python-docx processed {len(documents)} paragraphs")

            for i, doc in enumerate(documents, 1):
                print(f"\n  Paragraph {i}:")
                print(f"    Title: {doc.title}")
                print(f"    Content length: {len(doc.content)} chars")
                print(f"    Preview: {doc.content[:150]}...")

            return True
        except Exception as e:
            print(f"✗ python-docx backend test failed: {e}")
            import traceback
            traceback.print_exc()
            return False

    async def test_docx2txt_backend(self):
        """测试docx2txt后端"""
        print("\n" + "="*70)
        print("TEST 2: docx2txt Backend")
        print("="*70)

        processor = DocxProcessor()

        try:
            documents = await processor.process(str(self.test_docx), backend='docx2txt')
            print(f"✓ docx2txt processed {len(documents)} documents")

            for i, doc in enumerate(documents, 1):
                print(f"\n  Document {i}:")
                print(f"    Title: {doc.title}")
                print(f"    Content length: {len(doc.content)} chars")
                print(f"    Preview: {doc.content[:150]}...")

            return True
        except Exception as e:
            print(f"✗ docx2txt backend test failed: {e}")
            return False

    async def test_auto_backend_selection(self):
        """测试自动后端选择"""
        print("\n" + "="*70)
        print("TEST 3: Automatic Backend Selection")
        print("="*70)

        processor = DocxProcessor()

        try:
            documents = await processor.process(str(self.test_docx))
            print(f"✓ Auto-selected backend processed {len(documents)} paragraphs")

            for i, doc in enumerate(documents, 1):
                print(f"\n  Paragraph {i}:")
                print(f"    Title: {doc.title}")
                print(f"    Content length: {len(doc.content)} chars")

            return True
        except Exception as e:
            print(f"✗ Auto backend selection test failed: {e}")
            return False

    async def test_paragraph_extraction(self):
        """测试段落提取"""
        print("\n" + "="*70)
        print("TEST 4: Paragraph Extraction")
        print("="*70)

        processor = DocxProcessor()

        try:
            documents = await processor.process(str(self.test_docx))

            # 检查段落数量
            print(f"  Total paragraphs: {len(documents)}")

            # 检查标题段落
            title_docs = [doc for doc in documents if 'Chapter' in doc.content]
            print(f"  Title paragraphs: {len(title_docs)}")

            # 检查普通段落
            content_docs = [doc for doc in documents if doc.content not in [
                d.content for d in title_docs
            ]]
            print(f"  Content paragraphs: {len(content_docs)}")

            # 显示前3个段落
            print(f"\n  First 3 paragraphs:")
            for i, doc in enumerate(documents[:3], 1):
                print(f"\n    Paragraph {i}:")
                print(f"      Content: {doc.content[:100]}...")

            return len(documents) > 0
        except Exception as e:
            print(f"✗ Paragraph extraction test failed: {e}")
            return False

    async def test_table_extraction(self):
        """测试表格提取"""
        print("\n" + "="*70)
        print("TEST 5: Table Extraction")
        print("="*70)

        processor = DocxProcessor()

        try:
            documents = await processor.process(str(self.test_docx))

            # 查找包含表格内容的段落
            table_docs = []
            for doc in documents:
                if any(keyword in doc.content for keyword in [
                    'Feature', 'Status', 'Notes', 'Table'
                ]):
                    table_docs.append(doc)

            print(f"  Table-related paragraphs: {len(table_docs)}")

            if table_docs:
                print(f"\n  First table paragraph:")
                print(f"    Content: {table_docs[0].content[:200]}...")

            return len(table_docs) > 0
        except Exception as e:
            print(f"✗ Table extraction test failed: {e}")
            return False

    async def test_content_structure(self):
        """测试内容结构保留"""
        print("\n" + "="*70)
        print("TEST 6: Content Structure Preservation")
        print("="*70)

        processor = DocxProcessor()

        try:
            documents = await processor.process(str(self.test_docx))

            # 合并所有内容
            all_content = "\n".join(doc.content for doc in documents)

            # 检查预期结构元素
            expected_elements = [
                "DOCX Processor Test Document",
                "Chapter 1",
                "Section 1.1",
                "Section 1.2",
                "Chapter 2",
                "Chapter 3",
                "Implementation Steps",
                "Conclusion"
            ]

            found_elements = []
            for element in expected_elements:
                if element in all_content:
                    found_elements.append(element)
                    print(f"  ✓ Found: {element}")
                else:
                    print(f"  ✗ Missing: {element}")

            print(f"\n  Found {len(found_elements)}/{len(expected_elements)} structural elements")

            return len(found_elements) >= len(expected_elements) * 0.75
        except Exception as e:
            print(f"✗ Content structure test failed: {e}")
            return False

    async def test_metadata(self):
        """测试元数据完整性"""
        print("\n" + "="*70)
        print("TEST 7: Metadata Integrity")
        print("="*70)

        processor = DocxProcessor()

        try:
            documents = await processor.process(str(self.test_docx))

            if not documents:
                print("✗ No documents returned")
                return False

            doc = documents[0]
            required_fields = ['source_path', 'title', 'url']

            print(f"✓ Document has {len(doc.metadata)} metadata fields")

            for field in required_fields:
                if field in doc.metadata:
                    print(f"  ✓ {field}: {doc.metadata[field]}")
                else:
                    print(f"  ✗ Missing {field}")
                    return False

            # 检查DOCX特定字段
            docx_fields = ['paragraph_index', 'word_count']
            for field in docx_fields:
                if field in doc.metadata:
                    print(f"  ✓ {field}: {doc.metadata[field]}")

            return True
        except Exception as e:
            print(f"✗ Metadata test failed: {e}")
            return False

    async def test_chunking(self):
        """测试分块功能"""
        print("\n" + "="*70)
        print("TEST 8: Chunking")
        print("="*70)

        processor = DocxProcessor()

        try:
            # 测试段落分块
            documents = await processor.process(
                str(self.test_docx),
                chunk=True,
                chunk_size=1000,
                chunk_overlap=100
            )

            print(f"✓ Created {len(documents)} chunks")

            if documents:
                for i, doc in enumerate(documents[:3], 1):
                    print(f"\n  Chunk {i}:")
                    print(f"    Title: {doc.title}")
                    print(f"    Content length: {len(doc.content)} chars")
                    print(f"    Preview: {doc.content[:100]}...")

            return len(documents) > 0
        except Exception as e:
            print(f"✗ Chunking test failed: {e}")
            return False

    async def test_backend_fallback(self):
        """测试后端回退机制"""
        print("\n" + "="*70)
        print("TEST 9: Backend Fallback")
        print("="*70)

        processor = DocxProcessor()

        # 测试不存在的后端（应该回退到默认）
        try:
            documents = await processor.process(
                str(self.test_docx),
                backend='non-existent-backend'
            )
            print(f"✓ Fallback mechanism processed {len(documents)} paragraphs")

            return len(documents) > 0
        except Exception as e:
            print(f"✗ Backend fallback test failed: {e}")
            return False

    async def test_performance(self):
        """测试性能"""
        print("\n" + "="*70)
        print("TEST 10: Performance")
        print("="*70)

        import time

        processor = DocxProcessor()

        try:
            start_time = time.time()

            documents = await processor.process(str(self.test_docx))

            end_time = time.time()
            processing_time = end_time - start_time

            print(f"  Processing time: {processing_time:.3f} seconds")
            print(f"  Documents processed: {len(documents)}")
            print(f"  Average time per document: {processing_time/len(documents):.3f} seconds")

            if processing_time < 5.0:
                print(f"  ✓ Performance acceptable (< 5 seconds)")
                return True
            else:
                print(f"  ✗ Performance too slow (> 5 seconds)")
                return False

        except Exception as e:
            print(f"✗ Performance test failed: {e}")
            return False

    async def run_all_tests(self):
        """运行所有测试"""
        print("\n" + "="*70)
        print("DOCX PROCESSOR TEST SUITE")
        print("="*70)

        # 创建测试文件
        await self.create_test_docx()

        # 运行测试
        results = []
        results.append(("python-docx Backend", await self.test_python_docx_backend()))
        results.append(("docx2txt Backend", await self.test_docx2txt_backend()))
        results.append(("Auto Backend", await self.test_auto_backend_selection()))
        results.append(("Paragraph Extraction", await self.test_paragraph_extraction()))
        results.append(("Table Extraction", await self.test_table_extraction()))
        results.append(("Content Structure", await self.test_content_structure()))
        results.append(("Metadata", await self.test_metadata()))
        results.append(("Chunking", await self.test_chunking()))
        results.append(("Backend Fallback", await self.test_backend_fallback()))
        results.append(("Performance", await self.test_performance()))

        # 总结
        print("\n" + "="*70)
        print("TEST SUMMARY")
        print("="*70)

        passed = sum(1 for _, result in results if result)
        total = len(results)

        for test_name, result in results:
            status = "✓ PASS" if result else "✗ FAIL"
            print(f"  {status} - {test_name}")

        print(f"\n  Total: {passed}/{total} tests passed")

        return passed == total


async def main():
    """主函数"""
    test_suite = TestDocxProcessor()
    success = await test_suite.run_all_tests()

    if success:
        print("\n🎉 All tests passed!")
    else:
        print("\n⚠️  Some tests failed")

    print("\n" + "="*70)
    print("NOTES")
    print("="*70)
    print("• DOCX processor supports python-docx and docx2txt backends")
    print("• python-docx provides better structure preservation")
    print("• docx2txt is faster but may lose formatting")
    print("• Both backends automatically handle document structure")

    return success


if __name__ == "__main__":
    import logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    result = asyncio.run(main())
    sys.exit(0 if result else 1)
