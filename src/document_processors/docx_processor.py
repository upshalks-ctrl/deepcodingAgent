"""
DOCX处理器
支持Word文档
"""

import logging
from pathlib import Path
from typing import List

from src.models.document import Document
from src.document_processors import create_document_metadata

logger = logging.getLogger(__name__)


class DocxProcessor:
    """DOCX文档处理器"""

    async def process(self, file_path: str, **kwargs) -> List[Document]:
        """
        处理DOCX文档

        Args:
            file_path: 文件路径
            **kwargs: 额外参数

        Returns:
            List[Document]: 文档列表
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"DOCX文件不存在: {file_path}")

        logger.info(f"处理DOCX文件: {file_path}")

        # 方法1: 使用python-docx
        try:
            return await self._process_with_python_docx(file_path, **kwargs)
        except ImportError:
            logger.warning("python-docx未安装")
        except Exception as e:
            logger.error(f"python-docx处理失败: {e}")

        # 方法2: 使用docx2txt
        try:
            return await self._process_with_docx2txt(file_path, **kwargs)
        except ImportError:
            logger.warning("docx2txt未安装")
        except Exception as e:
            logger.error(f"docx2txt处理失败: {e}")

        raise RuntimeError("DOCX处理方法都失败了")

    async def _process_with_python_docx(self, file_path: str, **kwargs) -> List[Document]:
        """使用python-docx处理"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装python-docx: pip install python-docx")

        logger.info("使用python-docx处理")

        doc = DocxDocument(file_path)

        # 提取段落
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格
        tables = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                tables.append("\n".join(table_text))

        # 组合内容
        content_parts = []
        if paragraphs:
            content_parts.append("\n\n".join(paragraphs))
        if tables:
            content_parts.append("\n\n表格:\n" + "\n\n".join(tables))

        content = "\n\n".join(content_parts)

        # 创建元数据
        metadata = create_document_metadata(
            processor='python-docx',
            word_count=len(content),
            paragraph_count=len(paragraphs),
            table_count=len(tables)
        )

        document = Document(
            content=content,
            metadata=metadata,
            document_type='docx',
            source_path=file_path,
            title=Path(file_path).name,
            url = f"docx://{Path(file_path).name}"
        )

        logger.info(f"python-docx处理完成，共{len(content)}字符")
        return [document]

    async def _process_with_docx2txt(self, file_path: str, **kwargs) -> List[Document]:
        """使用docx2txt处理"""
        try:
            import docx2txt
        except ImportError:
            raise ImportError("请安装docx2txt: pip install docx2txt")

        logger.info("使用docx2txt处理")

        # 提取文本
        content = docx2txt.process(file_path)

        # 创建元数据
        metadata = create_document_metadata(
            processor='python-docx',
            word_count=len(content),
        )

        document = Document(
            content=content,
            metadata=metadata,
            document_type='docx',
            source_path=file_path,
            title=Path(file_path).name,
            url = f"docx://{Path(file_path).name}"
        )

        logger.info(f"docx2txt处理完成，共{len(content)}字符")
        return [document]
